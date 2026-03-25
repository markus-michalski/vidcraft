"""Document parsing for PDF, DOCX, and Markdown files.

Extracts structured text content, headings, and images from various
document formats to prepare content for video script generation.
"""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any


@dataclass
class DocumentSection:
    """A section extracted from a document."""

    heading: str
    level: int
    content: str
    word_count: int = 0
    has_images: bool = False
    has_code: bool = False
    has_list: bool = False

    def __post_init__(self) -> None:
        self.word_count = len(self.content.split())


@dataclass
class ParsedDocument:
    """Result of parsing a document."""

    title: str
    source_path: str
    format: str
    total_words: int
    sections: list[DocumentSection] = field(default_factory=list)
    metadata: dict[str, Any] = field(default_factory=dict)
    raw_text: str = ""

    def to_dict(self) -> dict[str, Any]:
        """Convert to serializable dict."""
        return {
            "title": self.title,
            "source_path": self.source_path,
            "format": self.format,
            "total_words": self.total_words,
            "section_count": len(self.sections),
            "sections": [
                {
                    "heading": s.heading,
                    "level": s.level,
                    "word_count": s.word_count,
                    "has_images": s.has_images,
                    "has_code": s.has_code,
                    "has_list": s.has_list,
                    "content": s.content,
                }
                for s in self.sections
            ],
            "metadata": self.metadata,
        }


def parse_document(path: Path) -> ParsedDocument:
    """Parse a document file into structured sections.

    Supports: .md, .pdf, .docx
    """
    suffix = path.suffix.lower()

    if suffix == ".md":
        return _parse_markdown(path)
    elif suffix == ".pdf":
        return _parse_pdf(path)
    elif suffix in (".docx", ".doc"):
        return _parse_docx(path)
    else:
        raise ValueError(f"Unsupported format: {suffix}. Supported: .md, .pdf, .docx")


def _parse_markdown(path: Path) -> ParsedDocument:
    """Parse a Markdown file into sections."""
    text = path.read_text(encoding="utf-8")

    # Strip YAML frontmatter
    text = re.sub(r"^---\s*\n.*?\n---\s*\n", "", text, flags=re.DOTALL)

    sections = _split_by_headings(text)
    title = sections[0].heading if sections else path.stem

    return ParsedDocument(
        title=title,
        source_path=str(path),
        format="markdown",
        total_words=len(text.split()),
        sections=sections,
        raw_text=text,
    )


def _parse_pdf(path: Path) -> ParsedDocument:
    """Parse a PDF file into sections."""
    try:
        import pdfplumber
    except ImportError:
        raise ImportError("pdfplumber required: pip install pdfplumber")

    all_text = []
    metadata: dict[str, Any] = {}

    with pdfplumber.open(path) as pdf:
        metadata["page_count"] = len(pdf.pages)
        if pdf.metadata:
            metadata["pdf_title"] = pdf.metadata.get("Title", "")
            metadata["pdf_author"] = pdf.metadata.get("Author", "")

        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                all_text.append(page_text)

    full_text = "\n\n".join(all_text)
    sections = _split_by_headings(full_text)

    # If no headings found, create one section per page
    if len(sections) <= 1 and len(all_text) > 1:
        sections = [
            DocumentSection(
                heading=f"Page {i + 1}",
                level=2,
                content=page_text,
            )
            for i, page_text in enumerate(all_text)
            if page_text.strip()
        ]

    title = metadata.get("pdf_title") or path.stem

    return ParsedDocument(
        title=title,
        source_path=str(path),
        format="pdf",
        total_words=len(full_text.split()),
        sections=sections,
        metadata=metadata,
        raw_text=full_text,
    )


def _parse_docx(path: Path) -> ParsedDocument:
    """Parse a DOCX file into sections."""
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx required: pip install python-docx")

    doc = Document(path)
    metadata: dict[str, Any] = {}

    if doc.core_properties:
        metadata["author"] = doc.core_properties.author or ""
        metadata["title"] = doc.core_properties.title or ""

    # Extract paragraphs grouped by headings
    sections: list[DocumentSection] = []
    current_heading = path.stem
    current_level = 1
    current_lines: list[str] = []
    has_images = False
    has_list = False

    for para in doc.paragraphs:
        style_name = (para.style.name or "").lower()

        # Detect headings
        if style_name.startswith("heading"):
            # Save previous section
            if current_lines:
                content = "\n".join(current_lines)
                sections.append(
                    DocumentSection(
                        heading=current_heading,
                        level=current_level,
                        content=content,
                        has_images=has_images,
                        has_list=has_list,
                    )
                )

            # Start new section
            current_heading = para.text.strip()
            try:
                current_level = int(style_name.replace("heading", "").strip())
            except ValueError:
                current_level = 2
            current_lines = []
            has_images = False
            has_list = False
        else:
            text = para.text.strip()
            if text:
                current_lines.append(text)
            if style_name.startswith("list"):
                has_list = True
            # Check for inline images
            if para.runs:
                for run in para.runs:
                    if run._element.findall(
                        ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing"
                    ):
                        has_images = True

    # Save last section
    if current_lines:
        content = "\n".join(current_lines)
        sections.append(
            DocumentSection(
                heading=current_heading,
                level=current_level,
                content=content,
                has_images=has_images,
                has_list=has_list,
            )
        )

    full_text = "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
    title = metadata.get("title") or path.stem

    return ParsedDocument(
        title=title,
        source_path=str(path),
        format="docx",
        total_words=len(full_text.split()),
        sections=sections,
        metadata=metadata,
        raw_text=full_text,
    )


def _split_by_headings(text: str) -> list[DocumentSection]:
    """Split text into sections by markdown-style headings."""
    lines = text.splitlines()
    sections: list[DocumentSection] = []
    current_heading = ""
    current_level = 1
    current_lines: list[str] = []

    for line in lines:
        heading_match = re.match(r"^(#{1,6})\s+(.+)$", line)
        if heading_match:
            # Save previous section
            if current_lines or current_heading:
                content = "\n".join(current_lines)
                sections.append(
                    DocumentSection(
                        heading=current_heading,
                        level=current_level,
                        content=content,
                        has_code="```" in content,
                        has_list=bool(re.search(r"^[-*\d]+[.)]\s", content, re.MULTILINE)),
                    )
                )

            current_heading = heading_match.group(2).strip()
            current_level = len(heading_match.group(1))
            current_lines = []
        else:
            current_lines.append(line)

    # Save last section
    if current_lines or current_heading:
        content = "\n".join(current_lines)
        sections.append(
            DocumentSection(
                heading=current_heading,
                level=current_level,
                content=content,
                has_code="```" in content,
                has_list=bool(re.search(r"^[-*\d]+[.)]\s", content, re.MULTILINE)),
            )
        )

    return sections


def extract_key_points(doc: ParsedDocument, max_points: int = 10) -> list[dict[str, str]]:
    """Extract key points from a parsed document.

    Identifies the most important content blocks based on:
    - Section headings (structural importance)
    - Lists (actionable items)
    - Code blocks (technical content needing visual support)
    - Short paragraphs after headings (topic sentences)
    """
    points: list[dict[str, str]] = []

    for section in doc.sections:
        if not section.content.strip():
            continue

        # Heading itself is a key point
        if section.heading:
            point: dict[str, str] = {
                "heading": section.heading,
                "type": "topic",
                "content": "",
                "video_relevance": "high" if section.has_code or section.has_images else "medium",
            }

            # Get first meaningful paragraph as summary
            paragraphs = [
                p.strip()
                for p in section.content.split("\n\n")
                if p.strip() and not p.strip().startswith("```")
            ]
            if paragraphs:
                point["content"] = paragraphs[0][:200]

            # Flag sections with visual content
            if section.has_code:
                point["type"] = "code_demo"
                point["note"] = "Contains code — consider screencast scene"
            elif section.has_list:
                point["type"] = "steps"
                point["note"] = "Contains list — consider step-by-step scenes"
            elif section.has_images:
                point["type"] = "visual"
                point["note"] = "Contains images — consider screenshot/demo scene"

            points.append(point)

            if len(points) >= max_points:
                break

    return points


def suggest_structure(
    doc: ParsedDocument, video_type: str = "tutorial"
) -> dict[str, Any]:
    """Suggest a video structure based on document content.

    Maps document sections to video scenes with timing estimates.
    """
    scenes: list[dict[str, Any]] = []
    total_words = 0

    # Add intro scene
    scenes.append(
        {
            "number": 1,
            "title": "Introduction",
            "type": "avatar",
            "source_section": "",
            "estimated_words": 40,
            "estimated_duration": "0:15",
            "notes": f"Introduce the topic: {doc.title}",
        }
    )
    total_words += 40

    # Map significant sections to scenes
    scene_num = 2
    for section in doc.sections:
        if not section.content.strip() or section.word_count < 20:
            continue

        # Determine visual type
        if section.has_code:
            visual_type = "screencast"
        elif section.has_images:
            visual_type = "split"
        elif section.has_list:
            visual_type = "slides"
        else:
            visual_type = "avatar"

        # Estimate narration words (compress document content)
        narration_words = min(section.word_count // 3, 150)
        narration_words = max(narration_words, 30)
        est_seconds = int(narration_words / 140 * 60)

        scenes.append(
            {
                "number": scene_num,
                "title": section.heading or f"Section {scene_num}",
                "type": visual_type,
                "source_section": section.heading,
                "estimated_words": narration_words,
                "estimated_duration": f"{est_seconds // 60}:{est_seconds % 60:02d}",
                "notes": section.content[:100].strip(),
            }
        )
        total_words += narration_words
        scene_num += 1

    # Add closing scene
    scenes.append(
        {
            "number": scene_num,
            "title": "Summary & Next Steps",
            "type": "avatar",
            "source_section": "",
            "estimated_words": 50,
            "estimated_duration": "0:20",
            "notes": "Recap key points and call to action",
        }
    )
    total_words += 50

    total_seconds = int(total_words / 140 * 60)

    return {
        "title": doc.title,
        "video_type": video_type,
        "scene_count": len(scenes),
        "estimated_total_words": total_words,
        "estimated_duration": f"{total_seconds // 60}:{total_seconds % 60:02d}",
        "scenes": scenes,
    }


def analyze_complexity(doc: ParsedDocument) -> dict[str, Any]:
    """Analyze document complexity to recommend video type and approach.

    Returns complexity metrics and recommendations.
    """
    code_sections = sum(1 for s in doc.sections if s.has_code)
    list_sections = sum(1 for s in doc.sections if s.has_list)
    total_sections = len(doc.sections)
    avg_section_words = doc.total_words / max(total_sections, 1)

    # Complexity score (0-100)
    complexity = 0
    complexity += min(doc.total_words / 50, 30)  # Length factor
    complexity += code_sections * 10  # Code factor
    complexity += min(total_sections * 3, 20)  # Depth factor
    complexity += 10 if avg_section_words > 200 else 0  # Dense sections
    complexity = min(complexity, 100)

    # Recommend video type
    if code_sections > total_sections * 0.5:
        recommended_type = "tutorial"
        reason = "High code content — step-by-step tutorial works best"
    elif list_sections > total_sections * 0.5:
        recommended_type = "installation-guide"
        reason = "Many sequential steps — installation guide format"
    elif doc.total_words < 500:
        recommended_type = "explainer"
        reason = "Short content — concise explainer video"
    elif doc.total_words > 2000:
        recommended_type = "training"
        reason = "Long content — structured training video or series"
    else:
        recommended_type = "product-demo"
        reason = "Mixed content — product demo format"

    # Estimate episode count
    if doc.total_words < 800:
        episodes = 1
    elif doc.total_words < 2000:
        episodes = 1
    elif doc.total_words < 5000:
        episodes = max(2, total_sections // 3)
    else:
        episodes = max(3, total_sections // 2)

    return {
        "complexity_score": round(complexity),
        "total_words": doc.total_words,
        "total_sections": total_sections,
        "code_sections": code_sections,
        "list_sections": list_sections,
        "avg_section_words": round(avg_section_words),
        "recommended_type": recommended_type,
        "recommended_reason": reason,
        "recommended_episodes": episodes,
    }
